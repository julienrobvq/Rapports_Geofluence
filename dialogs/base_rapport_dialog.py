# dialogs/base_rapport_dialog.py

from qgis.PyQt.QtWidgets import *
from qgis.core import QgsProject, QgsExpression
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Image
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph as RLParagraph, Spacer as RLSpacer, PageBreak as RLPageBreak, Image as RLImage
import os
from PyQt5.QtCore import QDate, QTime, QDateTime
from docx import Document
from docx.shared import Pt, Inches

class BaseRapportDialog(QDialog):

    """Classe générique réutilisable par tous les rapports"""

    def __init__(self, layer_form_name, champs_affiches, sections, parent=None, custom_mode=False):
        super().__init__(parent)
        
        self.custom_mode = custom_mode
        self.layer_form_name = layer_form_name
        self.champs_affiches = champs_affiches
        self.sections = sections

        #  Interface de base pour TOUS les rapports 
        self.setWindowTitle("Outil de création de rapport")
        self.resize(380, 420)
            

        layout = QVBoxLayout(self)
        layout.addWidget(QLabel("Sélectionnez un projet :"))
        self.proj_combo = QComboBox()
        layout.addWidget(self.proj_combo)

        self.lbl_titre = QLabel("Titre du rapport :")
        layout.addWidget(self.lbl_titre)
        self.titre_rapport = QLineEdit()
        layout.addWidget(self.titre_rapport)

        self.lbl_format = QLabel("Format d'exportation :")
        layout.addWidget(self.lbl_format)
        fmt_layout = QHBoxLayout()
        self.radio_pdf = QRadioButton("PDF")
        self.radio_word = QRadioButton("Word")
        self.radio_pdf.setChecked(True)
        fmt_layout.addWidget(self.radio_pdf)
        fmt_layout.addWidget(self.radio_word)
        layout.addLayout(fmt_layout)

        self.lbl_champs = QLabel("Champs à inclure dans le rapport :")
        layout.addWidget(self.lbl_champs)
        self.scroll = QScrollArea()
        container = QWidget()
        self.champ_layout = QGridLayout(container)
        self.scroll.setWidget(container)
        self.scroll.setWidgetResizable(True)
        layout.addWidget(self.scroll)

        # Boutons sélectionner / désélectionner
        btn_layout = QHBoxLayout()
        self.btn_select_all = QPushButton("Tout sélectionner")
        self.btn_unselect_all = QPushButton("Tout désélectionner")
        btn_layout.addWidget(self.btn_select_all)
        btn_layout.addWidget(self.btn_unselect_all)
        layout.addLayout(btn_layout)

        self.btn_select_all.clicked.connect(self.select_all)
        self.btn_unselect_all.clicked.connect(self.unselect_all)

        # Bouton OK
        self.btn_ok = QPushButton("Générer le rapport")
        layout.addWidget(self.btn_ok)
        self.btn_ok.clicked.connect(self.accept)

        #  Couches QGIS 
        self.layer_form_name = layer_form_name
        layers_form = QgsProject.instance().mapLayersByName(self.layer_form_name)
        if not layers_form:
            QMessageBox.critical(self, "Erreur", "Couche introuvable.")
            self._init_ok = False
            return
        
        self._init_ok = True
        self.layer_form = layers_form[0]

        # Couche événement commune à tous les rapports
        layers_evt = QgsProject.instance().mapLayersByName("Evenement")
        if not layers_evt:
            QMessageBox.critical(self, "Erreur", "Couche 'Evenement introuvable.")
            self.reject(); return

        self.layer_evt = layers_evt[0]

        # Ajout automatique des champs événement
        self.champs_evenement = ["Date", "Heure", "ID_Proj", "ID_Employ"]
        self.champs_affiches = self.champs_evenement + self.champs_affiches

        self.id_field_proj = "ID_Proj"

        self.remplir_projets()
        self.remplir_champs()

        # si le mode custom est activé, on masque certaines options

        if self.custom_mode:
            self.resize(380, 100)
            self.lbl_titre.hide()
            self.titre_rapport.hide()

            self.lbl_format.hide()
            self.radio_pdf.hide()
            self.radio_word.hide()

            self.lbl_champs.hide()
            self.scroll.hide()
            self.btn_select_all.hide()
            self.btn_unselect_all.hide()

    # Liste des projets
    def remplir_projets(self):
        field = self.layer_evt.fields().field(self.id_field_proj)
        cfg = field.editorWidgetSetup()

        projets_dict = {}

        for f in self.layer_evt.getFeatures():
            raw_value = f[self.id_field_proj]
            if raw_value in (None, "", " "):
                continue

            display_value = raw_value

            # ValueMap
            if cfg.type() == "ValueMap":
                mapping = cfg.config().get("map", {})
                display_value = mapping.get(str(raw_value), raw_value)

            # ValueRelation
            elif cfg.type() == "ValueRelation":
                rel_layer_id = cfg.config().get("Layer")
                key_field = cfg.config().get("Key")
                value_field = cfg.config().get("Value")
                rel_layer = QgsProject.instance().mapLayer(rel_layer_id)
                if rel_layer:
                    for rel_feat in rel_layer.getFeatures():
                        if str(rel_feat[key_field]) == str(raw_value):
                            display_value = rel_feat[value_field]
                            break

            projets_dict[str(raw_value)] = str(display_value)

        for raw, display in sorted(projets_dict.items(), key=lambda x: x[1]):
            self.proj_combo.addItem(display, raw)


    # Liste à cocher

    def remplir_champs(self):
        self.checkboxes = []
        for i, champ in enumerate(self.champs_affiches):

            if champ in self.layer_form.fields().names():
                alias = self.layer_form.fields().field(champ).alias() or champ

            elif champ in self.layer_evt.fields().names():
                alias = self.layer_evt.fields().field(champ).alias() or champ

            else:
                alias = champ

            cb = QCheckBox(alias)
            cb.setToolTip(champ)
            self.champ_layout.addWidget(cb, i, 0)
            self.checkboxes.append(cb)


    # Boutons
    def select_all(self):
        for cb in self.checkboxes:
            cb.setChecked(True)

    def unselect_all(self):
        for cb in self.checkboxes:
            cb.setChecked(False)

    # Section dialog
    def get_selection(self):
        id_proj = self.proj_combo.currentData()
        champs = [cb.toolTip() for cb in self.checkboxes if cb.isChecked()]
        titre = self.titre_rapport.text()
        return id_proj, champs, titre

    
    def get_display_value(self, layer, feature, field_name):
        return self._get_display_value_core(layer, feature, field_name)

    def get_evt_display_value(self, feat_evt, field_name):
        return self._get_display_value_core(self.layer_evt, feat_evt, field_name)

    # Génération du rapport 
    def _get_display_value_core(self, layer, feature, field_name):
        field = layer.fields().field(field_name)
        cfg = field.editorWidgetSetup()
        value = feature.attribute(field_name)
        
        # --- Choix multiples ---
        if isinstance(value, str) and value.startswith("{") and value.endswith("}"):
            raw = value[1:-1]
            items, current, in_quotes = [], "", False

            for c in raw:
                if c == '"':
                    in_quotes = not in_quotes
                elif c == "," and not in_quotes:
                    items.append(current.strip().strip('"'))
                    current = ""
                else:
                    current += c

            if current:
                items.append(current.strip().strip('"'))

            valeurs_affichees = []

            if cfg.type() == "ValueMap":
                raw_map = cfg.config().get("map", {})

                for v in items:
                    val_aff = v

                    if isinstance(raw_map, dict):
                        val_aff = raw_map.get(v, v)

                    elif isinstance(raw_map, list):
                        for item in raw_map:
                            if isinstance(item, dict):
                                for k, lbl in item.items():
                                    if str(k) == str(v):
                                        val_aff = str(lbl)

                    valeurs_affichees.append(val_aff)

            elif cfg.type() == "ValueRelation":
                rel_layer_id = cfg.config().get("Layer")
                key_field = cfg.config().get("Key")
                value_field = cfg.config().get("Value")
                rel_layer = QgsProject.instance().mapLayer(rel_layer_id)

                if rel_layer:
                    for v in items:
                        for f in rel_layer.getFeatures():
                            if str(f[key_field]) == str(v):
                                valeurs_affichees.append(str(f[value_field]))
                                break
            else:
                valeurs_affichees = items

            return ", ".join(valeurs_affichees)

        # --- Dates / heures ---
        if isinstance(value, QDateTime):
            return value.toString("yyyy-MM-dd HH:mm")
        if isinstance(value, QTime):
            return value.toString("HH:mm")
        if isinstance(value, QDate):
            return value.toString("yyyy-MM-dd")

        try:
            inner = value.toPyObject() if hasattr(value, "toPyObject") else value
            if isinstance(inner, QDateTime):
                return inner.toString("yyyy-MM-dd HH:mm")
            if isinstance(inner, QTime):
                return inner.toString("HH:mm")
            if isinstance(inner, QDate):
                return inner.toString("yyyy-MM-dd")
        except:
            pass

        if value in (None, ""):
            return ""

        # --- ValueMap ---
        if cfg.type() == "ValueMap":
            raw_map = cfg.config().get("map", {})

            # Cas 1 : ValueMap stocké sous forme de dict
            if isinstance(raw_map, dict):
                return raw_map.get(str(value), str(value))

            # Cas 2 : ValueMap stocké sous forme de liste
            elif isinstance(raw_map, list):
                for item in raw_map:
                    if isinstance(item, dict):
                        for k, v in item.items():
                            if str(k) == str(value):
                                return str(v)

            return str(value)

        # --- ValueRelation ---
        if cfg.type() == "ValueRelation":
            rel_layer_id = cfg.config().get("Layer")
            key_field = cfg.config().get("Key")
            value_field = cfg.config().get("Value")
            rel_layer = QgsProject.instance().mapLayer(rel_layer_id)

            if rel_layer:
                for f in rel_layer.getFeatures():
                    if str(f[key_field]) == str(value):
                        return str(f[value_field])

        return str(value)

    def accept(self):
        #  Récupération des paramètres 
        id_proj, champs, titre_rapport = self.get_selection()
        if self.custom_mode:
            fmt = "Word"
        else:
            fmt = "PDF" if self.radio_pdf.isChecked() else "Word"

        self.current_id_proj = id_proj
        self.current_champs = champs
        self.current_titre_rapport = titre_rapport

        #  Extensions suggérées 
        default_name = "Rapport.pdf" if fmt == "PDF" else "Rapport.docx"
        filter_str = (
            "PDF (*.pdf);;Tous les fichiers (*.*)"
            if fmt == "PDF"
            else "Word (*.docx);;Tous les fichiers (*.*)"
        )

        # Enregistrement
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Enregistrer le rapport",
            default_name,
            filter_str
        )

        if not file_path:
            return

        # Extension
        base, ext = os.path.splitext(file_path)
        ext = ext.lower()

        if fmt == "PDF" and ext != ".pdf":
            file_path = base + ".pdf"
        elif fmt == "Word" and ext != ".docx":
            file_path = base + ".docx"

        # Récupération des données

        layer_evt = self.layer_evt
        expr_evt = QgsExpression.createFieldEqualityExpression("ID_Proj", str(id_proj))
        layer_evt.selectByExpression(expr_evt)
        feats_evt = layer_evt.selectedFeatures()

        if not feats_evt:
            QMessageBox.warning(self, "Avertissement", f"Aucun événement trouvé pour {id_proj}.")
            return

        id_even_values = [
            f["ID_EVEN"] for f in feats_evt
            if f["ID_EVEN"] not in (None, "", " ")
        ]

        if not id_even_values:
            QMessageBox.warning(self, "Avertissement", "Aucun ID_EVEN trouvé pour ce projet.")
            return

        valeurs_str = ",".join([f"'{v}'" for v in id_even_values])
        expr_form = f'"ID_EVEN" IN ({valeurs_str})'
        self.layer_form.selectByExpression(expr_form)
        feats_form = self.layer_form.selectedFeatures()
        self.current_feats_evt = feats_evt
        self.current_feats_form = feats_form

        if not feats_form:
            QMessageBox.warning(self, "Avertissement", "Aucun enregistrement trouvé.")
            return

        # DCIM
        project_dir = os.path.dirname(QgsProject.instance().fileName())
        photo_root = os.path.join(project_dir, "DCIM")

        # Structure du rapport

        styles = getSampleStyleSheet()
        story = []

        # Titre 
        story.append(Paragraph(titre_rapport, styles["Title"]))
        story.append(Spacer(1, 12))

        #  Contenu
        for idx, feat in enumerate(feats_form):
            id_even = feat["ID_EVEN"]
            feat_evt = next((e for e in feats_evt if e["ID_EVEN"] == id_even), None)
            if not feat_evt:
                continue

            for titre, liste_champs in self.sections.items():
                contenu = []

                for champ in liste_champs:
                    if champ not in champs:
                        continue

                    # formulaire
                    if champ in self.layer_form.fields().names():
                        alias = self.layer_form.fields().field(champ).alias() or champ
                        valeur = self.get_display_value(self.layer_form, feat, champ)

                    # evenement
                    elif champ in self.layer_evt.fields().names():
                        alias = self.layer_evt.fields().field(champ).alias() or champ
                        valeur = self.get_evt_display_value(feat_evt, champ)

                    else:
                        continue

                    if valeur in ("", None, "NULL", "Null"):
                        continue

                    # photos
                    if "photo" in champ.lower():
                        rel_path = valeur.replace("/", os.sep).replace("\\", os.sep).lstrip(os.sep)
                        if rel_path.upper().startswith("DCIM" + os.sep.upper()):
                            rel_path = rel_path[len("DCIM" + os.sep):]
                        photo_path = os.path.normpath(os.path.join(photo_root, rel_path))
                        contenu.append(("photo", alias, photo_path))
                    else:
                        contenu.append(("texte", alias, valeur))

                if contenu:
                    story.append(Paragraph(titre, styles["Heading2"]))
                    story.append(Spacer(1, 6))

                    for typ, alias, valeur in contenu:
                        if typ == "texte":
                            story.append(Paragraph(f"<b>{alias}</b> : {valeur}", styles["BodyText"]))

                        elif typ == "photo":
                            if os.path.exists(valeur):
                                story.append(Spacer(1, 4))
                                img = Image(valeur)
                                img._restrictSize(A4[0] - 100, 300)
                                story.append(img)
                                story.append(Spacer(1, 4))

                    story.append(Spacer(1, 12))

            if idx < len(feats_form) - 1:
                story.append(PageBreak())

        # Exportation

        if fmt == "PDF":
            doc = SimpleDocTemplate(file_path, pagesize=A4)
            doc.build(story)
            QMessageBox.information(self, "Bravo", "PDF généré")
            super().accept()
            return

        # Word
        if fmt == "Word":
            self.export_word(file_path, story, titre_rapport)
            super().accept()
            return
    
    def exec_(self):
        if not getattr(self, "_init_ok", True):
            return 0
        return super().exec_()

    def export_word(self, file_path, story, titre_rapport):
        docx = Document()
        styles = docx.styles

        # Mise en page (ancien code)
        styles["Normal"].paragraph_format.space_after = Pt(1.5)
        styles["Normal"].paragraph_format.space_before = Pt(1.5)

        h1 = styles["Heading 1"]
        h1.paragraph_format.space_after = Pt(12)
        h1.paragraph_format.space_before = Pt(0)

        h2 = styles["Heading 2"]
        h2.paragraph_format.space_after = Pt(3)
        h2.paragraph_format.space_before = Pt(3)

        styles["List Bullet"].paragraph_format.space_after = Pt(3)
        styles["List Bullet"].paragraph_format.space_before = Pt(3)

        # Titre niveau 1
        docx.add_heading(titre_rapport or "Rapport", level=1)

        is_first = True  # évite qu'on ait deux fois le titre

        for elt in story:

            if is_first:
                is_first = False
                continue

            from reportlab.platypus import Paragraph as RLParagraph, Spacer as RLSpacer, PageBreak as RLPageBreak, Image as RLImage

            if isinstance(elt, RLParagraph):
                txt = elt.text.replace("<b>", "").replace("</b>", "")

                if elt.style.name == "Heading2":
                    docx.add_heading(txt, level=2)
                else:
                    if " : " in txt:
                        alias, valeur = txt.split(" : ", 1)
                        p = docx.add_paragraph(style="List Bullet")
                        run_alias = p.add_run(f"{alias} : ")
                        run_alias.bold = True
                        p.add_run(valeur)
                    else:
                        docx.add_paragraph(txt, style="List Bullet")

            elif isinstance(elt, RLPageBreak):
                docx.add_page_break()

            elif isinstance(elt, RLImage):
                img_path = elt.filename
                if os.path.exists(img_path):
                    docx.add_picture(img_path, width=Inches(2.2))

        docx.save(file_path)
        QMessageBox.information(self, "Bravo", "Word généré")

        super().accept()